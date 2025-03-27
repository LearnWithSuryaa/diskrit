from docx import Document
from docx.shared import Pt

# Kunci RSA tetap
e = 79
d = 1019
n = 3337

def name_to_ascii(name):
    ascii_values = [ord(char) for char in name]

    # Menampilkan tabel ASCII
    output = []
    table_data = []
    output.append("=" * 40)
    output.append("| No | Char  | ASCII |")
    output.append("=" * 40)
    for i, char in enumerate(name, 1):
        ascii_val = ord(char)
        output.append(f"| {i:<2} | {char:<5} | {ascii_val:<5} |")
        table_data.append((str(i), char, str(ascii_val)))
    output.append("=" * 40)

    # Menggabungkan semua ASCII menjadi satu string
    ascii_str = "".join(str(ascii_val) for ascii_val in ascii_values)
    output.append(f"\nKetika digabung menjadi satu, menjadi :")
    output.append(ascii_str)

    # Memecah ke dalam blok 3 digit
    ascii_blocks = [ascii_str[i:i+3] for i in range(0, len(ascii_str), 3)]
    output.append("\nJika dipecah menjadi blok m dengan ukuran masing-masing 3 digit:")
    output.append(" ".join(ascii_blocks))

    return [int(block) for block in ascii_blocks if block.isdigit()], "\n".join(output), table_data, ascii_blocks

def rsa_encrypt(m, index):
    output = []
    output.append(f"Enkripsi Blok {index}:\n")
    output.append(f"c{index} = {m}^{e} mod {n}\n")

    result = pow(m, e, n)
    steps = []
    powers = {1: m % n}

    exponent_order = [1, 2, 4, 8, 16, 32, 64, 72, 76, 78, 79]
    base_mapping = {72: (8, 64), 76: (4, 72), 78: (2, 76), 79: (1, 78)}

    for exp in exponent_order:
        if exp == 1:
            temp_result = m % n
            step_text = f"{m}^{exp} mod {n} = {temp_result}"
        elif exp in {2, 4, 8, 16, 32, 64}:
            base_exp = exp // 2
            temp_result = (powers[base_exp] * powers[base_exp]) % n
            step_text = (
                f"{m}^{exp} mod {n} = [{m}^{base_exp} mod {n}] . [{m}^{base_exp} mod {n}] mod {n}\n"
                f"               = [{powers[base_exp]}] . [{powers[base_exp]}] mod {n}\n"
                f"               = {powers[base_exp] * powers[base_exp]} mod {n}\n"
                f"               = {temp_result}"
            )
        else:
            base1, base2 = base_mapping[exp]
            temp_result = (powers[base1] * powers[base2]) % n
            step_text = (
                f"{m}^{exp} mod {n} = [{m}^{base1} mod {n}] . [{m}^{base2} mod {n}] mod {n}\n"
                f"               = [{powers[base1]}] . [{powers[base2]}] mod {n}\n"
                f"               = {powers[base1] * powers[base2]} mod {n}\n"
                f"               = {temp_result}"
            )
        
        powers[exp] = temp_result
        steps.append(step_text)
        output.append(step_text)

    output.append(f"Hasil akhir enkripsi: c{index} = {result}\n")
    return result, "\n".join(output)

def rsa_decrypt(c, index):
    output = []
    output.append(f"Dekripsi Blok {index}:\n")
    output.append(f"m{index} = {c}^{d} mod {n}\n")

    result = pow(c, d, n)
    steps = []
    powers = {1: c % n}
    
    output.append(f"{c}^1 mod {n} = {powers[1]}")

    # Urutan eksponen yang ditampilkan
    displayed_exponents = [1, 2, 4, 8, 16, 32, 64, 128, 256, 512, 768, 896, 960, 992, 1008, 1016, 1018, 1019]

    for exp in sorted(displayed_exponents):
        if exp == 1:
            temp_result = c % n
            step_text = f"{c}^{exp} mod {n} = {temp_result}"
        elif exp in powers:
            continue
        else:
            base_exp = exp // 2
            if base_exp in powers:
                temp_result = (powers[base_exp] * powers[base_exp]) % n
                step_text = (
                    f"{c}^{exp} mod {n} = [{c}^{base_exp} mod {n}] . [{c}^{base_exp} mod {n}] mod {n}\n"
                    f"               = [{powers[base_exp]}] . [{powers[base_exp]}] mod {n}\n"
                    f"               = {powers[base_exp] * powers[base_exp]} mod {n}\n"
                    f"               = {temp_result}"
                )
            else:
                base1 = max([x for x in powers if x < exp])
                base2 = exp - base1
                temp_result = (powers[base1] * powers[base2]) % n
                step_text = (
                    f"{c}^{exp} mod {n} = [{c}^{base1} mod {n}] . [{c}^{base2} mod {n}] mod {n}\n"
                    f"               = [{powers[base1]}] . [{powers[base2]}] mod {n}\n"
                    f"               = {powers[base1] * powers[base2]} mod {n}\n"
                    f"               = {temp_result}"
                )
        
        powers[exp] = temp_result
        steps.append(step_text)
        output.append(step_text)

    output.append(f"Hasil akhir dekripsi: m{index} = {result}\n")
    return result, "\n".join(output)

def save_to_docx(name, table_data, ascii_blocks, encryption_steps, decryption_steps):
    doc = Document()

    # Tambahkan judul
    doc.add_heading("Laporan Enkripsi dan Dekripsi RSA", level=1)

    # Informasi Kunci RSA
    doc.add_heading("Informasi Kunci RSA", level=2)
    doc.add_paragraph(f"Modulus (n): {n}")
    doc.add_paragraph(f"Eksponen Publik (e): {e}")
    doc.add_paragraph(f"Eksponen Privat (d): {d}")

    # Konversi ASCII
    doc.add_heading("Konversi Nama ke ASCII", level=2)
    doc.add_paragraph(f"Nama yang dienkripsi: {name}")
    doc.add_paragraph("Tabel ASCII:")

    table = doc.add_table(rows=len(table_data) + 1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text = "No", "Char", "ASCII"

    for i, (num, char, ascii_val) in enumerate(table_data):
        row_cells = table.rows[i + 1].cells
        row_cells[0].text, row_cells[1].text, row_cells[2].text = num, char, ascii_val

    doc.add_paragraph(f"Blok ASCII: {' '.join(ascii_blocks)}")

    # Hasil Enkripsi
    doc.add_heading("Hasil Enkripsi RSA", level=2)
    for step in encryption_steps:
        doc.add_paragraph(step)

    # Hasil Dekripsi
    doc.add_heading("Hasil Dekripsi RSA", level=2)
    for step in decryption_steps:
        doc.add_paragraph(step)

    doc.save("[NIM-Klean]_DISKRIT [Kelas-Klean]_Tugas3.docx")
    print("\nProses selesai! Hasil telah disimpan di '[NIM-Klean]_DISKRIT [Kelas-Klean]_Tugas3.docx'.")

def main():
    name = input("Masukkan nama: ")
    ascii_blocks, ascii_output, table_data, ascii_block_list = name_to_ascii(name)
    
    encryption_steps = []
    decryption_steps = []

    for i, block in enumerate(ascii_blocks, start=1):
        encrypted, encrypt_text = rsa_encrypt(block, i)
        encryption_steps.append(encrypt_text)

        decrypted, decrypt_text = rsa_decrypt(encrypted, i)
        decryption_steps.append(decrypt_text)

    save_to_docx(name, table_data, ascii_block_list, encryption_steps, decryption_steps)

if __name__ == "__main__":
    main()
