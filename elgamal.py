import random
from tabulate import tabulate
from docx import Document

# Parameter kunci publik
p = 2579
alpha = 2
beta = 949

# Fungsi untuk menghitung perpangkatan modular dengan tampilan langkah-langkah
def mod_exp_verbose(base, exp, mod):
    steps = []
    intermediate_values = {1: base % mod}

    # Langkah awal (pangkat 1)
    steps.append(f"{base}¹    mod {mod:<4} = {intermediate_values[1]}")

    # Hitung kelipatan pangkat 2 (2, 4, 8, 16, ...)
    power = 2
    while power <= exp:
        prev = intermediate_values[power // 2]
        intermediate_values[power] = (prev * prev) % mod
        steps.append(f"{base}^{power:<4} mod {mod:<4} = [{base}^{power//2:<2} mod {mod:<4}] . [{base}^{power//2:<2} mod {mod:<4}] mod {mod}")
        steps.append(f"                      = [{prev:<4}] . [{prev:<4}] mod {mod}")
        steps.append(f"                      = {prev * prev:<6} mod {mod}")
        steps.append(f"                      = {intermediate_values[power]}\n")
        power *= 2

    # Jika exp tidak ditemukan langsung di dictionary, cari kombinasi pangkat
    if exp not in intermediate_values:
        bin_exp = bin(exp)[2:]  # Representasi biner dari exp
        exp_parts = sorted([2**i for i in range(len(bin_exp)) if bin_exp[-(i+1)] == '1'], reverse=True)

        # Format baru untuk menampilkan kombinasi pangkat
        steps.append(f"Menghitung kombinasi pangkat terbesar yang mendekati:")
        steps.append(f"     {base}^{exp:<4} mod {mod:<4} = " + " . ".join([f"[{base}^{part:<2} mod {mod:<4}]" for part in exp_parts]) + f" mod {mod}")

        # Hitung hasil dengan format langkah per langkah
        result = 1
        for i, part in enumerate(exp_parts):
            if i == 0:
                result = intermediate_values[part]
                steps.append(f"                      = {result}")
            else:
                prev_result = result
                result = (result * intermediate_values[part]) % mod
                steps.append(f"                      = [{prev_result:<4}] . [{intermediate_values[part]:<4}] mod {mod}")
                steps.append(f"                      = {prev_result * intermediate_values[part]:<6} mod {mod}")
                steps.append(f"                      = {result}\n")
        intermediate_values[exp] = result
        
    return intermediate_values[exp], steps

# Fungsi enkripsi dengan algoritma El Gamal
def elgamal_encrypt(message, p, alpha, beta):
    encrypted_pairs = []
    output_text = ""

    # Data untuk tabel
    table_data = [["i", "Karakter", "Plainteks mi", "ASCII"]]

    # Daftar angka k yang unik
    available_k_values = [x for x in range(10, 100, 2)]  

    for i, char in enumerate(message):
        m = ord(char)  # Konversi karakter ke ASCII

        # Pilih nilai k yang belum digunakan
        k = random.choice(available_k_values)
        available_k_values.remove(k)

        table_data.append([i + 1, char, f"m{i+1}", m])

        output_text += f"\n{i+1}.   m{i+1} = {m}\n"
        output_text += f"     k{i+1}  = {k}\n"

        # Hitung gamma = α^k mod p
        gamma, gamma_steps = mod_exp_verbose(alpha, k, p)
        output_text += f"     γ{i+1}  = {alpha}^{k} mod {p}\n"
        output_text += "     " + "\n     ".join(gamma_steps) + "\n"
        output_text += f"     Ditemukan nilai ɣ{i+1} = {gamma}\n\n"

        # Hitung delta = β^k mod p
        delta_base, delta_steps = mod_exp_verbose(beta, k, p)
        output_text += f"     δ{i+1}  = {beta}^{k} mod {p}\n"
        output_text += "     " + "\n     ".join(delta_steps) + "\n"
        output_text += f"     Ditemukan nilai δ{i+1} = {delta_base}\n\n"

        output_text += f"     Menghitung rumus δ{i+1} = {beta}^{k} . m{i+1} mod {p}\n"
        output_text += f"                  = [{beta}^{k} mod {p}] . [{m}] mod {p}\n"
        output_text += f"                  = [{delta_base}] . [{m}] mod {p}\n"
        output_text += f"                  = {delta_base * m} mod {p}\n"

        # Hitung δi = (β^k * m) mod p
        delta = (delta_base * m) % p
        output_text += f"                  = {delta}\n"
        output_text += f"     Ditemukan nilai δ{i+1} = {delta}\n\n"

        encrypted_pairs.append((gamma, delta))

        output_text += f"     Hasil enkripsi m{i+1} = {m} dengan k{i+1} = {k} adalah ({gamma}, {delta})\n"

    return encrypted_pairs, output_text, table_data

# Fungsi untuk menyimpan hasil enkripsi ke dalam file DOCX
def save_to_docx(filename, output_text, table_data):
    doc = Document()
    
    # Tambahkan judul
    doc.add_heading("Hasil Enkripsi El Gamal", level=1)
    
    # Tambahkan tabel
    doc.add_paragraph("Kunci Publik (p = 2579 ,  α = 2 , β = 949)")
    doc.add_paragraph("Konversi karakter pesan ke kode ASCII:")
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Table Grid'  # Menggunakan gaya tabel grid

    for row_idx, row_data in enumerate(table_data):
        for col_idx, col_data in enumerate(row_data):
            table.cell(row_idx, col_idx).text = str(col_data)

    doc.add_paragraph("\n")  # Tambahkan spasi

    # Tambahkan hasil perhitungan
    doc.add_paragraph(output_text)

    # Simpan ke file
    doc.save(filename)

# Main program
if __name__ == "__main__":
    message = input("Masukkan pesan yang ingin dienkripsi: ")
    encrypted_data, output_text, table_data = elgamal_encrypt(message, p, alpha, beta)

    # Simpan hasil ke file DOCX
    output_filename = "[NIM-Klean]_DISKRIT [Kelas-Klean]_Tugas4.docx"
    save_to_docx(output_filename, output_text, table_data)
    print(f"\nHasil enkripsi berhasil disimpan ke dalam file {output_filename}")
